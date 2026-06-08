from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Any
import io

def to_docx(plan: dict) -> bytes:
    doc = Document()
    pid = plan.get('project_id','')
    t = doc.add_heading(f'Project {pid}', level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def af(label: str, value: Any):
        if not value: return
        p = doc.add_paragraph()
        r = p.add_run(f'{label}: '); r.bold = True; r.font.size = Pt(10)
        p.add_run(str(value)).font.size = Pt(10)

    mi = plan.get('main_image',{})
    if mi:
        doc.add_heading('1. 主图方案', level=1)
        af('目标', mi.get('goal')); af('构图', mi.get('composition')); af('背景', mi.get('background'))
        af('光影', mi.get('lighting')); af('文案', mi.get('copywriting'))
        if mi.get('prompt'):
            p = doc.add_paragraph(mi['prompt']); p.style = 'Body Text'
            for run in p.runs: run.font.size = Pt(9); run.font.color.rgb = RGBColor(0x33,0x99,0x33)

    wb = plan.get('white_bg',{})
    if wb:
        doc.add_heading('2. 白底图方案', level=1)
        af('目标', wb.get('goal')); af('指令', wb.get('instructions'))

    scenes = plan.get('scene_images',[])
    if scenes:
        doc.add_heading('3. 场景图方案', level=1)
        for i,s in enumerate(scenes,1):
            doc.add_heading(f'场景{i}: {s.get("scene_name","")}', level=2)
            af('用户', s.get('target_user')); af('叙事', s.get('scene_narrative'))
            if s.get('prompt'):
                p = doc.add_paragraph(s['prompt']); p.style = 'Body Text'
                for run in p.runs: run.font.size = Pt(9); run.font.color.rgb = RGBColor(0x33,0x99,0x33)

    sps = plan.get('selling_points',[])
    if sps:
        doc.add_heading('4. 卖点图模块', level=1)
        for i,sp in enumerate(sps,1):
            doc.add_heading(f'{i}. {sp.get("title","")}', level=2)
            af('描述', sp.get('description')); af('视觉', sp.get('visual_representation'))

    scripts = plan.get('video_scripts',[])
    if scripts:
        doc.add_heading('5. 视频脚本', level=1)
        for s in scripts:
            doc.add_heading(f'{s.get("duration_seconds","?")}秒', level=2)
            af('目标', s.get('video_goal')); af('CTA', s.get('cta'))
            sb = s.get('storyboard',[])
            if sb:
                tbl = doc.add_table(rows=1, cols=5); tbl.style = 'Table Grid'
                for j,h in enumerate(['镜头','时长','画面','字幕','旁白']): tbl.rows[0].cells[j].text = h
                for shot in sb:
                    row = tbl.add_row()
                    row.cells[0].text = str(shot.get('shot_number',''))
                    row.cells[1].text = str(shot.get('duration',''))
                    row.cells[2].text = str(shot.get('visual',''))
                    row.cells[3].text = str(shot.get('subtitle',''))
                    row.cells[4].text = str(shot.get('voiceover',''))

    ad = plan.get('ad_material',{})
    if ad:
        doc.add_heading('6. 广告素材方案', level=1)
        af('目标', ad.get('ad_goal')); af('人群', ad.get('target_audience'))
        af('角度', ad.get('ad_angle')); af('CTA', ad.get('cta'))

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read()
